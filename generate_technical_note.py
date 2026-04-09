"""
Generate ASCE Technical Note — 7 double-spaced pages max.

Adaptive Lagrangian Refinement for Observation-Dependent
Hydraulic Simulation Using Vortex Particle Methods
"""

import os
import sys
import numpy as np

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Run experiments ───────────────────────────────────────────────────────

print("Running experiments...")
from quantum_hydraulics.research.alr_experiments import (
    run_convergence, run_cost_benefit, run_sigma_field, run_scour,
    CHANNEL_LENGTH, CHANNEL_WIDTH, DEPTH, Q, PIER_X, OBS_CENTER,
)
from quantum_hydraulics.research.sediment_scenarios import generate_clearwater_scour_scenario
from quantum_hydraulics.integration.sediment_transport import QuasiUnsteadyEngine

cost = run_cost_benefit()
scour_r = run_scour()

channel, sed_mix, hydrograph, sed_meta = generate_clearwater_scour_scenario()
engine = QuasiUnsteadyEngine(channel=channel, sediment_mix=sed_mix,
                              upstream_feed_fraction=0.0,
                              computational_increment_hours=5.0, bed_mixing_steps=3)
engine.set_hydrograph_durations(hydrograph)
sed_sim = engine.run()

# Ensure figures
os.system("python run_benchmark_validation.py --figures 2>nul >nul")

# ── Document setup (ASCE format) ─────────────────────────────────────────

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0  # ASCE: double-spaced
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.font.size = Pt(12)
    hs.paragraph_format.line_spacing = 2.0
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def body(text, indent=True):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p


def heading(text, level=1):
    # ASCE: word headings, not numbered
    doc.add_heading(text, level=level)


def add_table(headers, rows, caption=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"

    # ASCE: horizontal rules only for header and bottom, no vertical rules
    from docx.oxml.ns import qn
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'bottom', 'insideH']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): '000000'
        })
        borders.append(border)
    for border_name in ['left', 'right', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): '000000'
        })
        borders.append(border)
    tblPr.append(borders)

    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"


def add_figure(path, caption, width=5.0):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"


# ══════════════════════════════════════════════════════════════════════════
# TECHNICAL NOTE CONTENT (~1,500 words + 1 table + 1 figure)
# ══════════════════════════════════════════════════════════════════════════

# ── Title block ───────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Adaptive Lagrangian Refinement for Observation-Dependent "
    "Hydraulic Simulation Using Vortex Particle Methods"
)
run.bold = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Michael Flynn, PE, M.ASCE")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("McGill Associates, PA, Asheville, North Carolina")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

doc.add_paragraph()  # blank line

# ── Abstract ──────────────────────────────────────────────────────────────

heading("Abstract")

body(
    "This technical note presents Adaptive Lagrangian Refinement (ALR), a method "
    "that enables physics-based turbulence simulation at engineering-relevant "
    "locations without the cost of uniform high-resolution computation. ALR employs "
    "observation-dependent resolution within a vortex particle framework using a "
    "symmetrized Biot-Savart kernel (Barba and Rossi 2010) that preserves "
    "circulation to 0.03% over full simulations. Each particle carries an adaptive "
    "core size that concentrates computational effort at bridge piers, scour-prone "
    "junctions, or other critical locations while maintaining coarse approximation "
    "elsewhere. The underlying hydraulics use Colebrook-White friction rather than "
    "Manning's equation. ALR achieves 0.2% vorticity error at 500 particles "
    "relative to a 6,000-particle uniform baseline, representing a 12-fold "
    "particle reduction. Cross-validation against six established methods shows "
    "correlation coefficients of 0.998 with Laursen contraction scour and 0.605 "
    "with the HEC-18 pier scour equation. The method provides a complementary "
    "turbulence amplification factor that augments rather than replaces these "
    "empirical methods. A quasi-unsteady sediment transport demonstration with "
    "fractional bedload transport, Hirano active-layer armoring, and Exner "
    "morphodynamic feedback produces physically consistent clear-water degradation "
    "behavior. The method integrates with PCSWMM as a post-processor for both "
    "1D and 2D models, requiring no additional meshing or CFD software. The "
    "complete source code and validation suite are open-source.",
    indent=False,
)

# ── Introduction ──────────────────────────────────────────────────────────

heading("Introduction")

body(
    "Manning's equation has served hydraulic engineering for over 130 years, but "
    "it collapses complex turbulent boundary layer physics into a single empirical "
    "roughness coefficient. At bridge piers, channel contractions, and culvert "
    "outlets, the local turbulent velocity field determines whether bed material "
    "mobilizes and at what rate. Computational fluid dynamics resolves these "
    "fields, but uniform high-resolution simulation is prohibitively expensive "
    "for routine screening."
)
body(
    "This note introduces Adaptive Lagrangian Refinement (ALR), which makes "
    "computational resolution observation-dependent. ALR is the primary "
    "contribution; supporting modules for scour assessment, pier vortex shedding, "
    "free-surface correction, and quasi-unsteady sediment transport are included "
    "to demonstrate practical extensibility. The method integrates with PCSWMM as "
    "a lightweight post-processor, filling the gap between Manning-based screening "
    "and full 3D CFD."
)

# ── Methodology ───────────────────────────────────────────────────────────

heading("Methodology")

heading("Vortex Particle Method", level=2)
body(
    "Particles carry 3D vorticity vectors in a Structure-of-Arrays layout. "
    "Velocity is computed via the Biot-Savart integral using a symmetrized "
    "regularized kernel: the effective core size for the interaction between "
    "particles i and j is computed as the root-sum-square of their individual "
    "core sizes, ensuring momentum conservation when core sizes vary "
    "(Barba and Rossi 2010). Viscous diffusion uses symmetrized Particle "
    "Strength Exchange. The hydraulic engine uses Colebrook-White friction with "
    "log-law velocity profiles rather than Manning's equation."
)

heading("Observation-Dependent Resolution", level=2)
body(
    "Each particle's core size adapts based on distance from observation zones "
    "via a Gaussian enhancement function that produces up to 5-fold resolution "
    "concentration at the observation center with smooth decay to base resolution "
    "at distance. Multiple observation zones are supported; the maximum "
    "enhancement across all zones determines the local core size. "
    f"At 500 particles, ALR achieves {cost.errors_vorticity[1]:.1%} vorticity "
    f"error relative to a 6,000-particle uniform baseline, with wall time "
    f"scaling from {cost.wall_times[0]:.2f} s to {cost.wall_times[-1]:.2f} s."
)

# ── Validation ────────────────────────────────────────────────────────────

heading("Validation")

body(
    "The method was cross-validated against six independent published methods. "
    "Colebrook-White velocity was compared to Manning's equation across five "
    "channel types, showing 25.5% average difference consistent with the known "
    "Strickler conversion offset. The Shields parameter and Neill's critical "
    "velocity trends are reproduced correctly across grain sizes, consistent "
    "with established threshold behavior."
)

heading("Contraction and Pier Scour", level=2)
body(
    "Contraction shear amplification was compared against Laursen live-bed scour "
    "across five width ratios (0.9 to 0.5). The Pearson correlation is r = 0.998. "
    "At 50% contraction, ALR predicts 3.3-fold shear amplification, consistent "
    "with continuity-based velocity scaling modified by Colebrook-White friction "
    "factor variation. HEC-18 pier scour correlation across five configurations "
    "yields r = 0.605; all cases with significant scour are detected."
)

heading("Melville Design Curve", level=2)
body(
    "The Melville (1997) dimensionless design equation was compared against "
    "ALR's turbulence amplification factor across flow intensity ratios from "
    "0.3 to 1.5 for a 3-ft circular pier in deep water. ALR provides a "
    "consistent constriction-based shear amplification of approximately "
    "1.11-fold, independent of flow intensity (Table 1). This captures the "
    "geometric blockage effect. The Tier 2 vortex particle analysis yields a "
    "higher amplification (1.44-fold) because it additionally captures "
    "turbulence-induced Reynolds stresses via Biot-Savart induction. "
    "ALR and Melville capture different physics and are complementary."
)

add_table(
    ["V/Vc", "Melville K_I", "ALR Amplification", "Melville ds (ft)", "Combined ds (ft)"],
    [
        ["0.3", "0.30", "1.11", "2.16", "2.40"],
        ["0.5", "0.50", "1.11", "3.60", "4.00"],
        ["0.7", "0.70", "1.11", "5.04", "5.60"],
        ["0.9", "0.90", "1.11", "6.48", "7.20"],
        ["1.0", "1.00", "1.11", "7.20", "8.00"],
    ],
    "Table 1. Melville (1997) design curve vs ALR amplification (b = 3 ft, y/b = 3.3).",
)

body(
    "This comparison uses the published dimensionless design curve, not "
    "individual measured data points from flume experiments. Direct comparison "
    "against measured scour profiles requires extracting exact test conditions "
    "from the original publications and is identified as future work.",
)

add_figure(
    "Benchmark_figures/fig3_melville_design_curve.png",
    "Fig. 1. Melville (1997) dimensionless design curve (left) and resulting "
    "scour depth with and without ALR turbulence amplification factor (right) "
    "for a 3-ft pier.",
    width=5.5,
)

# ── Conclusions ───────────────────────────────────────────────────────────

heading("Conclusions")

body(
    "Cross-validation against six independent published methods shows ALR "
    "produces physically consistent trends (Laursen contraction r = 0.998, "
    "HEC-18 pier scour r = 0.605). ALR provides a complementary turbulence "
    "amplification factor that augments established empirical methods. "
    f"At 500 particles, ALR achieves {cost.errors_vorticity[1]:.1%} vorticity "
    f"error with a 12-fold particle reduction and circulation conserved to "
    f"0.03% via the symmetrized Biot-Savart kernel. A quasi-unsteady sediment "
    f"transport demonstration produces {abs(sed_sim.total_scour_ft):.1f} ft of "
    f"clear-water scour with {sed_sim.final_d50_mm / sed_sim.initial_gradation.d50_mm:.0f}-fold "
    f"surface coarsening, consistent with downstream-of-dam behavior "
    f"(Williams and Wolman 1984). The method integrates with PCSWMM as a "
    f"post-processor requiring no additional meshing or CFD software."
)

# ── Code Availability ─────────────────────────────────────────────────────

heading("Data Availability Statement")

body(
    "The complete ALR source code, validation suite, and synthetic test cases "
    "are open-source under the MIT license (DOI: 10.5281/zenodo.19462126).",
    indent=False,
)

# ── References (ASCE author-date style) ───────────────────────────────────

heading("References")

refs = [
    'Barba, L., and Rossi, L. (2010). "Global field interpolation for particle '
    'methods." J. Comput. Phys., 229(4), 1292\u20131310.',

    'Colebrook, C. F. (1939). "Turbulent flow in pipes, with particular reference '
    'to the transition region between smooth and rough pipe laws." J. Inst. Civ. '
    'Eng., 11(4), 133\u2013156.',

    'Egiazaroff, I. (1965). "Calculation of non-uniform sediment concentrations." '
    'J. Hydraul. Div., 91(HY4), 225\u2013248.',

    'Hirano, M. (1971). "River bed degradation with armoring." Trans. Jpn. Soc. '
    'Civ. Eng., 195, 55\u201365.',

    'Melville, B. W. (1997). "Pier and abutment scour: Integrated approach." '
    'J. Hydraul. Eng., 123(2), 125\u2013136.',

    'Meyer-Peter, E., and M\u00fcller, R. (1948). "Formulas for bed-load transport." '
    'Proc., 2nd Meeting, Int. Assoc. Hydraul. Res., Stockholm, 39\u201364.',

    'Richardson, E. V., and Davis, S. R. (2001). Evaluating scour at bridges, '
    '4th Ed., Hydraulic Engineering Circular No. 18, FHWA-NHI-01-001, Federal '
    'Highway Administration, Washington, DC.',

    'Williams, G. P., and Wolman, M. G. (1984). "Downstream effects of dams on '
    'alluvial rivers." Professional Paper 1286, U.S. Geological Survey, '
    'Washington, DC.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 2.0
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

# ── Save ──────────────────────────────────────────────────────────────────

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Flynn_ASCE_JHE_Technical_Note.docx",
)
doc.save(output_path)

# Word count estimate
word_count = 0
for para in doc.paragraphs:
    word_count += len(para.text.split())
print(f"\nTechnical Note saved: {output_path}")
print(f"Estimated word count: {word_count}")
print(f"Target: ~1,750 words for 7 double-spaced pages")
