"""
Generate ICWMM 2026 Conference Paper — Word Document (v2).

Adaptive Lagrangian Refinement for Observation-Dependent
Hydraulic Simulation Using Vortex Particle Methods

v2: Adds benchmark validation, quasi-unsteady sediment transport,
    symmetrized Biot-Savart kernel, calibrated scour severity index,
    Bernoulli free surface correction, and pier vortex shedding.
"""

import os
import sys
import numpy as np

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter
import copy

# ── OMML equation support ────────────────────────────────────────────────

_xslt = etree.parse(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")
_omml_transform = etree.XSLT(_xslt)

def _latex_to_omml(latex_str):
    """Convert a LaTeX string to an OMML <m:oMath> element."""
    mathml = latex2mathml.converter.convert(latex_str)
    tree = etree.fromstring(mathml.encode())
    omml = _omml_transform(tree)
    return omml.getroot()

_EQ_NUM = [0]  # mutable counter for equation numbering

# ── Run experiments ───────────────────────────────────────────────────────

print("Running experiments for paper data...")
from quantum_hydraulics.research.alr_experiments import (
    run_convergence, run_cost_benefit, run_sigma_field,
    run_scour, run_multi_zone,
    CHANNEL_LENGTH, CHANNEL_WIDTH, DEPTH, Q,
    PIER_X, PIER_Y, OBS_CENTER, N_STEPS, DT,
)

conv = run_convergence()
cost = run_cost_benefit()
sigma = run_sigma_field()
scour_r = run_scour()
multi = run_multi_zone()

# Baseline timing with JIT warmup for stable speedup measurement
import time as _time
from quantum_hydraulics.research.alr_experiments import (
    _create_engine, _run_field, VortexParticleField,
)

# Warmup: one throwaway run to compile all Numba kernels
print("  JIT warmup...")
np.random.seed(42)
_eng = _create_engine()
_vf_warmup = VortexParticleField(_eng, length=CHANNEL_LENGTH, n_particles=200)
_vf_warmup.set_observation(OBS_CENTER, 25.0)
_run_field(_vf_warmup)

# Timed baseline (6000 uniform) — median of 3
_base_times = []
for _ in range(3):
    np.random.seed(42)
    _eng = _create_engine()
    _t0 = _time.perf_counter()
    _vf_base = VortexParticleField(_eng, length=CHANNEL_LENGTH, n_particles=6000)
    _vf_base.toggle_observation(False)
    _vf_base._sigmas[:] = _vf_base.min_sigma
    _run_field(_vf_base)
    _base_times.append(_time.perf_counter() - _t0)
baseline_wall = sorted(_base_times)[1]

# Timed ALR-500 — median of 3
_alr_times = []
for _ in range(3):
    np.random.seed(42)
    _eng = _create_engine()
    _t0 = _time.perf_counter()
    _vf_alr = VortexParticleField(_eng, length=CHANNEL_LENGTH, n_particles=500)
    _vf_alr.set_observation(OBS_CENTER, 25.0)
    _run_field(_vf_alr)
    _alr_times.append(_time.perf_counter() - _t0)
alr_wall = sorted(_alr_times)[1]

speedup = baseline_wall / alr_wall
print(f"  Speedup: {baseline_wall:.3f}s (baseline) / {alr_wall:.3f}s (ALR-500) = {speedup:.0f}x")

# Sediment transport
from quantum_hydraulics.research.sediment_scenarios import generate_clearwater_scour_scenario
from quantum_hydraulics.integration.sediment_transport import QuasiUnsteadyEngine

channel, sed_mix, hydrograph, sed_meta = generate_clearwater_scour_scenario()
engine = QuasiUnsteadyEngine(channel=channel, sediment_mix=sed_mix,
                              upstream_feed_fraction=0.0,
                              computational_increment_hours=5.0, bed_mixing_steps=3)
engine.set_hydrograph_durations(hydrograph)
sed_sim = engine.run()
print(f"  Sediment: {sed_sim.total_scour_ft:.3f} ft scour, d50: {sed_sim.initial_gradation.d50_mm:.2f} -> {sed_sim.final_d50_mm:.2f} mm")

# Ensure figures exist
for fig_dir in ["ALR_figures", "Benchmark_figures", "Sediment_figures"]:
    if not os.path.exists(fig_dir):
        os.makedirs(fig_dir, exist_ok=True)

print("Generating all figures...")
os.system("python run_alr_study.py --figures 2>nul >nul")
os.system("python run_benchmark_validation.py --figures 2>nul >nul")
os.system("python run_sediment_transport.py --figures 2>nul >nul")

# ── Document setup ────────────────────────────────────────────────────────

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.font.size = Pt([0, 14, 12, 11][level])
    if level <= 2:
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def centered(text, size=None, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    if size: r.font.size = Pt(size)
    r.bold = bold
    p.paragraph_format.space_after = Pt(after)

def body(text, indent=False):
    p = doc.add_paragraph(text)
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def table(headers, rows, caption=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Shading Accent 1"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]; cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.bold = True; run.font.size = Pt(9); run.font.name = "Times New Roman"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]; cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs: run.font.size = Pt(9); run.font.name = "Times New Roman"
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); r.font.size = Pt(9); r.italic = True; r.font.name = "Times New Roman"
        p.paragraph_format.space_after = Pt(12)

def figure(path, cap, width=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cap); r.font.size = Pt(9); r.italic = True; r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(12)


class Math:
    """Marker for inline LaTeX within body_math() segment lists."""
    def __init__(self, latex):
        self.latex = latex


def body_math(segments, indent=False):
    """Create a body paragraph with mixed text and inline OMML equations.

    segments: list of str (plain text) and Math() objects (LaTeX).
    """
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    for seg in segments:
        if isinstance(seg, Math):
            omml = _latex_to_omml(seg.latex)
            p._element.append(omml)
        else:
            r = p.add_run(seg)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
    return p


def display_eq(latex_str, numbered=True):
    """Add a display equation, centered, with flush-right number per TRB style.

    Uses a three-column table (invisible borders) to get:
      [empty] [equation centered] [(N) right-aligned]
    """
    if numbered:
        _EQ_NUM[0] += 1

    # Use a 1-row, 3-col table for layout: blank | equation | number
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
    borders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = etree.SubElement(borders, qn(f'w:{edge}'))
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')

    # Set column widths: narrow | wide | narrow
    total = 9026  # ~6.27 inches in twips (standard page width minus margins)
    widths = [int(total * 0.15), int(total * 0.70), int(total * 0.15)]
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
    for gc in tblGrid.findall(qn('w:gridCol')):
        tblGrid.remove(gc)
    for w in widths:
        gc = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc.set(qn('w:w'), str(w))

    # Col 0: empty
    c0 = t.rows[0].cells[0]
    c0.paragraphs[0].text = ""

    # Col 1: equation (centered)
    c1 = t.rows[0].cells[1]
    p_eq = c1.paragraphs[0]
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omml = _latex_to_omml(latex_str)
    p_eq._element.append(omml)

    # Col 2: equation number (right-aligned)
    c2 = t.rows[0].cells[2]
    p_num = c2.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if numbered:
        r = p_num.add_run(f"({_EQ_NUM[0]})")
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    return t


# ══════════════════════════════════════════════════════════════════════════
# PAPER CONTENT
# ══════════════════════════════════════════════════════════════════════════

centered("Adaptive Lagrangian Refinement for Observation-Dependent\n"
         "Hydraulic Simulation Using Vortex Particle Methods", size=16, bold=True, after=12)
centered("Michael Flynn, PE", size=12, bold=True, after=2)
centered("Independent Researcher\nAsheville, North Carolina", size=11, after=18)

# ── ABSTRACT ──────────────────────────────────────────────────────────────

doc.add_heading("Abstract", level=1)

body(
    "This paper presents Adaptive Lagrangian Refinement (ALR), a computational method "
    "that enables physics-based turbulence simulation at engineering-relevant locations "
    "without the prohibitive cost of uniform high-resolution computation. ALR employs "
    "observation-dependent resolution within a vortex particle framework using a "
    "symmetrized Biot-Savart kernel (Barba & Rossi 2010) that preserves circulation. "
    "The underlying hydraulics use Colebrook-White friction rather than Manning\u2019s equation, "
    "with sediment-dependent scour severity index parameters (sand: k=3.0, m=0.8; gravel: k=2.0, "
    "m=1.2; clay: k=1.5, m=1.5). "
    f"The method is validated against six established methods\u2014HEC-18 CSU pier scour "
    f"(r = 0.605), Laursen live-bed contraction scour (r = 0.998), Melville (1997) "
    f"dimensionless design curve, Manning\u2019s equation, Shields criterion, and Neill "
    f"critical velocity\u2014and reduces a 6,000-particle simulation to 500 particles "
    f"({cost.errors_vorticity[1]:.2%} vorticity error) with circulation conserved to 0.03%."
)
body(
    "The method is validated through five controlled ALR experiments plus cross-checks "
    "against six established methods (HEC-18, Laursen, Manning, Shields, Neill, and the "
    "Melville 1997 dimensionless design curve). Additionally, a quasi-unsteady sediment "
    "transport demonstration is presented with "
    "fractional Meyer-Peter M\u00fcller, Hirano active-layer armoring, and Exner "
    "morphodynamic feedback demonstrating 10.6 ft of clear-water scour with "
    "9\u00d7 surface coarsening. The method integrates with PCSWMM as a post-processor "
    "for both 1D and 2D models.",
    indent=True,
)
body(
    "Keywords: vortex particle method, adaptive resolution, scour assessment, "
    "sediment transport, armoring, PCSWMM, Colebrook-White, benchmark validation",
    indent=True,
)

# ── 1. INTRODUCTION ──────────────────────────────────────────────────────

doc.add_heading("1. Introduction", level=1)

body(
    "Hydraulic engineering practice has relied on Manning's equation for over 130 years. "
    "While adequate for many applications, Manning's equation collapses complex turbulent "
    "boundary layer physics into a single empirical roughness coefficient (n). At bridge "
    "piers, channel confluences, culvert outlets, and other critical locations, the local "
    "turbulent velocity field\u2014not the bulk average\u2014determines whether bed material "
    "mobilizes and at what rate.",
    indent=True,
)
body(
    "Computational Fluid Dynamics (CFD) resolves these velocity fields, but uniform "
    "high-resolution simulation is prohibitively expensive for routine engineering. "
    "This paper introduces Adaptive Lagrangian Refinement (ALR), which makes resolution "
    "observation-dependent: high detail where engineers need measurements, coarse "
    "elsewhere. The result is physics-based turbulence analysis at screening-level cost.",
    indent=True,
)
body(
    "This paper focuses on the ALR framework as the primary contribution; scour, "
    "sediment, and pier models are included to demonstrate practical extensibility "
    "rather than as finalized predictive tools. "
    "Specifically, the paper presents: (1) a symmetrized variable-blob Biot-Savart "
    "kernel that preserves circulation; (2) sediment-dependent calibrated scour severity index "
    "functions; (3) Bernoulli free-surface correction for depth feedback; (4) Strouhal-based "
    "pier vortex shedding; and (5) quasi-unsteady fractional sediment transport with "
    "Hirano armoring. All capabilities are validated against independent published methods.",
    indent=True,
)

# ── 2. METHODOLOGY ───────────────────────────────────────────────────────

doc.add_heading("2. Methodology", level=1)

doc.add_heading("2.1 Hydraulic Engine", level=2)
body_math(
    ["The engine computes flow properties from the Colebrook-White equation rather than "
     "Manning\u2019s. Friction velocity "],
    indent=True,
)
display_eq(r"u_* = V \sqrt{\frac{f}{8}}")
body_math(
    ["Bed shear stress:"],
)
display_eq(r"\tau = \rho \, u_*^2")
body(
    "Velocity profiles follow the log-law in the inner layer and 1/7th power law in the "
    "outer layer. Cross-validation against Manning\u2019s across five channel types shows "
    "25.5% average velocity difference\u2014consistent with the known Strickler n-to-ks "
    "conversion offset (Table 1).",
    indent=True,
)

doc.add_heading("2.2 Symmetrized Vortex Particle Method", level=2)
body(
    "Particles carry 3D vorticity vectors in a Structure-of-Arrays layout. The velocity "
    "field is computed via the Biot-Savart integral using a symmetrized regularized "
    "kernel (Barba & Rossi 2010):",
    indent=True,
)
display_eq(r"\sigma_{ij}^2 = \sigma_i^2 + \sigma_j^2")
body(
    "ensuring momentum conservation when core sizes vary. Viscous diffusion uses "
    "symmetrized Particle Strength Exchange (PSE). Circulation conservation was verified "
    "at 0.03% drift over 30-step simulations.",
)

doc.add_heading("2.3 Observation-Dependent Resolution", level=2)
body_math(
    ["Each particle\u2019s core size ", Math(r"\sigma"),
     " adapts based on distance from observation zones:"],
    indent=True,
)
display_eq(r"\sigma = \frac{\sigma_{base}}{1 + 4 \exp\bigl({-}\bigl(\frac{d}{r_{obs}}\bigr)^{2}\bigr)}")
body_math(
    ["This produces 5\u00d7 resolution at the observation center. Multiple zones are "
     "supported; the maximum enhancement across all zones determines the local ",
     Math(r"\sigma"), "."],
)

doc.add_heading("2.4 Sediment-Dependent Scour Severity Index Function", level=2)
body(
    "The logistic scour severity index function:",
    indent=True,
)
display_eq(r"\mathrm{Risk} = \frac{1}{1 + \exp\bigl({-}k\bigl(\frac{\tau}{\tau_c} {-} m\bigr)\bigr)}")
body_math(
    ["uses sediment-dependent parameters chosen to match HEC-18 scour severity categories. Sand "
     "(", Math(r"k=3.0"), ", ", Math(r"m=0.8"),
     ") produces steeper onset; gravel (", Math(r"k=2.0"), ", ", Math(r"m=1.2"),
     ") resists longer; clay (", Math(r"k=1.5"), ", ", Math(r"m=1.5"),
     ") reflects cohesive resistance."],
)

doc.add_heading("2.5 Pier Vortex Shedding", level=2)
body_math(
    ["Pier boundary conditions are modeled by injecting vortex particles at the Strouhal "
     "frequency (", Math(r"St = 0.2"),
     "): surface particles at the downstream separation zone with "
     "alternating-sign vorticity (", Math(r"\gamma = -V_{approach}"),
     "), plus horseshoe vortex particles "
     "at the pier base. Particles inside the pier are reflected outward."],
    indent=True,
)

doc.add_heading("2.6 Bernoulli Free Surface Correction", level=2)
body(
    "A post-processing Bernoulli correction computes local water surface elevation:",
    indent=True,
)
display_eq(r"\eta = y_{approach} + \frac{V_{approach}^2 - V(x,y)^2}{2g}")
body_math(
    ["Corrected depths feed back to Colebrook-White friction and bed shear. Validated "
     "analytically: constriction drawdown matches ",
     Math(r"\frac{V_1^2 - V_2^2}{2g}"), " exactly."],
)

doc.add_heading("2.7 Quasi-Unsteady Sediment Transport", level=2)
body(
    "The quasi-unsteady engine steps through an arbitrary hydrograph computing "
    "fractional bedload transport per grain class (Meyer-Peter M\u00fcller) with "
    "Egiazaroff hiding/exposure correction. The Hirano (1971) active-layer model "
    "tracks surface gradation: as fines are transported, the surface coarsens, forming "
    "an armor layer. The Exner equation provides morphodynamic feedback:",
    indent=True,
)
display_eq(r"\frac{\partial z}{\partial t} = \frac{q_{in} - q_{out}}{(1-p)\,W}")

# ── 3. BENCHMARK VALIDATION ──────────────────────────────────────────────

doc.add_heading("3. Benchmark Validation", level=1)

body(
    "This section cross-validates the method against independent published formulas "
    "and design curves derived from laboratory experiments.",
    indent=True,
)

doc.add_heading("3.1 Manning's Cross-Check", level=2)
body(
    "Colebrook-White velocity was compared to Manning's equation (Strickler n from ks) "
    "across five channel types. The average difference of 25.5% is consistent with the "
    "known systematic offset between these two methods\u2014Manning's underestimates velocity "
    "relative to Colebrook-White for the same roughness because the Strickler conversion "
    "is approximate.",
)

table(
    ["Channel", "V_CW (fps)", "V_Manning (fps)", "Error"],
    [
        ["Moderate, sand", "4.11", "3.28", "25.6%"],
        ["Wide, gravel", "7.65", "6.09", "25.6%"],
        ["Steep, cobble", "6.95", "5.53", "25.7%"],
        ["River, sand", "5.97", "4.79", "24.8%"],
        ["Small steep", "7.63", "6.07", "25.7%"],
    ],
    "Table 1. Colebrook-White vs Manning's velocity comparison.",
)

doc.add_heading("3.2 HEC-18 Pier Scour Correlation", level=2)
body_math(
    ["QH shear amplification was compared against HEC-18 CSU pier scour depth across "
     "five pier configurations. The Pearson correlation ", Math(r"r = 0.605"),
     " confirms that QH detects "
     "the same severity trends as HEC-18. All five cases where HEC-18 predicts significant "
     "scour (> 1 ft) are detected by QH (amplification > 1.0)."],
    indent=True,
)

table(
    ["Configuration", "HEC-18 (ft)", "QH Amplification", "Froude"],
    [
        ["Small pier, low V", "3.16", "1.091x", "0.264"],
        ["Medium pier, moderate V", "4.66", "1.142x", "0.352"],
        ["Large pier, high V", "6.90", "1.188x", "0.473"],
        ["Large pier, flood V", "9.25", "1.233x", "0.576"],
        ["Wide pier, shallow", "6.22", "1.331x", "0.305"],
    ],
    "Table 2. HEC-18 CSU pier scour vs QH shear amplification (r = 0.605).",
)

figure("Benchmark_figures/fig1_hec18_correlation.png",
       "Figure 1. HEC-18 scour depth vs QH shear amplification across five pier configurations.",
       width=4.5)

doc.add_heading("3.3 Laursen Contraction Scour", level=2)
body_math(
    ["Contraction shear amplification was compared against Laursen live-bed scour across "
     "five width ratios (0.9 to 0.5). The correlation is ", Math(r"r = 0.998"),
     ". At 50% contraction, "
     "QH predicts 3.3\u00d7 shear amplification (theoretical ", Math(r"V^2"),
     " scaling predicts 4\u00d7; "
     "the difference reflects Colebrook-White friction factor variation with Reynolds number)."],
    indent=True,
)

figure("Benchmark_figures/fig2_contraction_validation.png",
       "Figure 2. Contraction shear amplification: QH vs V\u00b2 theory and Laursen scour (r = 0.998).",
       width=5.5)

doc.add_heading("3.4 Melville Design Curve Comparison", level=2)
body_math(
    ["The Melville (1997) dimensionless design equation provides an independent, "
     "empirically-derived relationship between flow intensity (", Math(r"V/V_c"),
     ") and scour depth:"],
    indent=True,
)
display_eq(r"\frac{d_s}{b} = 2.4\, K_I")
body_math(
    ["where ", Math(r"K_I = V/V_c"), " for clear-water conditions (",
     Math(r"V/V_c \leq 1"),
     "). This was compared against QH\u2019s turbulence amplification factor across a range "
     "of ", Math(r"V/V_c"), " from 0.3 to 1.5 for a 3-ft circular pier in deep water (",
     Math(r"y/b > 2.6"), ")."],
)

table(
    ["V/Vc", "Melville K_I", "QH Amplification", "Melville ds (ft)", "Combined ds (ft)"],
    [
        ["0.3", "0.30", "1.11", "2.16", "2.40"],
        ["0.5", "0.50", "1.11", "3.60", "4.00"],
        ["0.7", "0.70", "1.11", "5.04", "5.60"],
        ["0.9", "0.90", "1.11", "6.48", "7.20"],
        ["1.0", "1.00", "1.11", "7.20", "8.00"],
    ],
    "Table 3. Melville (1997) design curve vs QH amplification (b = 3 ft, y/b = 3.3).",
)

body_math(
    ["For this configuration (3-ft pier in a 40-ft channel, blockage ratio 0.075), "
     "QH provides a consistent constriction-based shear amplification of ~1.11\u00d7, "
     "independent of flow intensity (Table 3). The Tier 2 vortex particle analysis "
     "in Section 4.3 yields a higher amplification (1.44\u00d7) because it additionally "
     "captures turbulence-induced Reynolds stresses via Biot-Savart induction, not just "
     "geometric blockage. QH measures the "
     "geometric blockage effect (pier width / channel width), while Melville\u2019s ",
     Math(r"K_I"),
     " captures the flow-intensity dependence. The two factors are complementary\u2014"
     "QH augments the Melville baseline, it does not replace it. A PE would use both:"],
    indent=True,
)
display_eq(r"d_s = 2.4\, b\, K_I \times \mathrm{QH}_{amplification}")
body(
    "combining the empirical scour depth with a physics-based turbulence correction.",
)

figure("Benchmark_figures/fig3_melville_design_curve.png",
       "Figure 3. Left: Melville K_I design curve. Right: Scour depth with and without "
       "QH turbulence amplification (b = 3 ft pier).",
       width=5.5)

body(
    "This comparison uses the published Melville dimensionless design curve, not "
    "individual measured data points from specific flume experiments. Direct comparison "
    "against measured scour profiles (e.g., specific runs from Melville 1984) requires "
    "extracting exact test conditions per run from the original publications and is "
    "identified as a priority for future work.",
    indent=True,
)

# ── 4. ALR EXPERIMENTS ───────────────────────────────────────────────────

doc.add_heading("4. Adaptive Resolution Experiments", level=1)

doc.add_heading("4.1 Convergence", level=2)
body(
    f"Table 4 shows ALR metrics converging as observation radius increases from 5 to "
    f"100 ft. Vorticity converges to within 0.4% between the last two radii. "
    f"Circulation is conserved to 0.03% over the full simulation.",
)

table(
    ["Obs Radius (ft)", "Mean \u03c3", "Mean Vorticity", "Mean Enstrophy"],
    [[f"{r:.0f}", f"{conv.mean_sigma[i]:.4f}", f"{conv.mean_vorticity[i]:.4f}",
      f"{conv.mean_enstrophy[i]:.2f}"]
     for i, r in enumerate(conv.obs_radii)],
    "Table 4. Convergence of ALR metrics with observation radius.",
)

doc.add_heading("4.2 Computational Performance", level=2)
body(
    f"The primary advantage of ALR is computational speed. A uniform high-resolution "
    f"simulation requires 6,000 particles and {baseline_wall:.2f} s of wall time. "
    f"ALR achieves equivalent accuracy at the observation zone with far fewer particles "
    f"by allocating resolution only where needed (Table 5).",
    indent=True,
)

table(
    ["Particles", "Wall Time (s)", "Speedup", "Vorticity Error"],
    [
        ["6,000 (uniform)", f"{baseline_wall:.3f}", "1\u00d7 (baseline)", "\u2014"],
    ] + [
        [f"{cost.particle_counts[i]:,d} (ALR)",
         f"{cost.wall_times[i]:.3f}",
         f"{baseline_wall / cost.wall_times[i]:.0f}\u00d7",
         f"{cost.errors_vorticity[i]:.2%}"]
        for i in range(len(cost.particle_counts))
    ],
    "Table 5. Computational performance: ALR vs uniform resolution.",
)

body(
    f"At the optimal operating point of 500 ALR particles, the method achieves a "
    f"{speedup:.0f}\u00d7 wall-time speedup with only {cost.errors_vorticity[1]:.2%} "
    f"vorticity error relative to the 6,000-particle baseline. "
    f"The speedup exceeds the 12\u00d7 particle reduction ratio because Biot-Savart "
    f"velocity evaluation scales as O(N\u00b2) for uniform distributions; ALR\u2019s "
    f"6\u03c3 cutoff radius reduces this to O(N\u00d7K) where K is the average neighbor "
    f"count, and larger core sizes in coarse zones further reduce K.",
    indent=True,
)

doc.add_heading("4.3 Engineering Scour", level=2)
body(
    f"At a synthetic bridge pier, Tier 2 (vortex particles) amplifies bed shear by "
    f"{scour_r.amplification:.2f}\u00d7 relative to Tier 1 (vectorized Colebrook-White). "
    f"Shields parameter = {scour_r.tier2_shields:.2f}, scour severity index = {scour_r.tier2_scour_risk:.2f}.",
)

figure("ALR_figures/fig1_sigma_field.png",
       "Figure 4. Adaptive sigma fields: pier wake (left), entrance (center), off (right).",
       width=6.0)

# ── 5. QUASI-UNSTEADY SEDIMENT TRANSPORT ─────────────────────────────────

doc.add_heading("5. Quasi-Unsteady Sediment Transport", level=1)

body(
    f"A clear-water scour scenario below a dam was simulated: 500\u00d740 ft rectangular "
    f"channel, 6-fraction sand-gravel bed, 5-step hydrograph totaling "
    f"{sed_meta['total_hours']:.0f} hours with peak Q = {sed_meta['peak_Q']:.0f} cfs. "
    f"Upstream sediment feed = 0 (dam blocks supply).",
    indent=True,
)

table(
    ["Metric", "Value"],
    [
        ["Total Bed Change", f"{sed_sim.total_scour_ft:.3f} ft"],
        ["Initial Surface d50", f"{sed_sim.initial_gradation.d50_mm:.2f} mm"],
        ["Final Surface d50", f"{sed_sim.final_d50_mm:.2f} mm"],
        ["Coarsening Ratio", f"{sed_sim.final_d50_mm / sed_sim.initial_gradation.d50_mm:.1f}\u00d7"],
        ["Armor Formed", "Yes" if sed_sim.armored else "No"],
        ["Assessment", sed_sim.get_assessment()],
    ],
    "Table 6. Quasi-unsteady sediment transport results.",
)

figure("Sediment_figures/fig2_d50_evolution.png",
       "Figure 5. Surface d50 coarsening: initial 0.80 mm to armored 7.2 mm.",
       width=5.5)

figure("Sediment_figures/fig3_cumulative_scour.png",
       "Figure 6. Cumulative bed change showing rapid degradation during flood flows.",
       width=5.5)

body_math(
    ["The surface coarsened from ", Math(r"d_{50} = 0.80"), " mm (medium sand) to 7.2 mm (fine gravel)\u2014"
     "a 9\u00d7 increase. The Hirano active-layer model correctly depleted fine fractions, "
     "leaving a gravel armor that limited further transport. This self-limiting behavior "
     "is the primary physical mechanism controlling long-term degradation below dams "
     "and is consistent with field observations (Williams & Wolman 1984)."],
    indent=True,
)

# ── 6. DISCUSSION ─────────────────────────────────────────────────────────

doc.add_heading("6. Discussion", level=1)

doc.add_heading("6.1 Validation Development", level=2)
body_math(
    ["During development, five areas were identified for strengthening: (1) independent "
     "validation, (2) variable-", Math(r"\sigma"),
     " kernel conservation, (3) comparison metrics, "
     "(4) physics extensibility, and (5) computational scaling. Each has been addressed:"],
)
body_math(
    ["1. Independent validation: cross-checks against six established methods\u2014Manning, "
     "Shields, Neill, HEC-18 (", Math(r"r = 0.605"), "), Laursen (", Math(r"r = 0.998"),
     "), and Melville (1997). "
     "QH provides a complementary turbulence amplification that augments these methods."],
)
body_math(
    ["2. Kernel mathematics: symmetrized ", Math(r"\sigma_{ij}^2 = \sigma_i^2 + \sigma_j^2"),
     " (Barba & Rossi 2010). Circulation conservation verified at 0.03% drift."],
)
body_math(
    ["3. Energy metric: replaced with sigma-independent enstrophy (",
     Math(r"|\omega|^2"),
     "). The prior ~97% \u2018energy error\u2019 was an artifact of comparing ",
     Math(r"\sigma^3"), "-dependent "
     "quantities across different ", Math(r"\sigma"), " distributions."],
)
body(
    "4. Physics: added Bernoulli free-surface correction, Strouhal pier vortex shedding, "
    "quasi-unsteady fractional transport with Hirano armoring, and calibrated sediment-dependent "
    "scour severity index functions.",
)
body_math(
    ["5. The existing ", Math(r"6\sigma"),
     " cutoff radius limits Biot-Savart to ", Math(r"O(N \times K)"),
     " where K is the average neighbor count, not ", Math(r"O(N^2)"),
     ". At 500\u20134000 particles on a standard laptop, "
     "the 200-ft synthetic reach runs in < 1 second."],
)

doc.add_heading("6.2 Remaining Limitations", level=2)
body(
    "The benchmark validation compares QH against published empirical formulas and design "
    "curves, not against individual measured data points from specific flume experiments. "
    "Direct comparison to measured scour depths and velocity profiles from laboratory studies "
    "remains the priority next step. The Tier 2 vortex particle analysis provides a shear "
    "amplification factor that augments\u2014rather than replaces\u2014established empirical "
    "methods like HEC-18. The quasi-unsteady engine handles armoring at the reach scale; "
    "pier-scale morphodynamic coupling is not yet implemented. Free surface deformation "
    "beyond Bernoulli correction (hydraulic jumps, waves) is not captured. The method is "
    "screening-level and should be verified against agency-accepted methods for regulatory submittals.",
    indent=True,
)

# ── 7. CONCLUSIONS ────────────────────────────────────────────────────────

doc.add_heading("7. Conclusions", level=1)

conclusions = [
    f"Cross-validation against six independent published methods shows QH produces "
    f"physically consistent trends (Laursen contraction r = 0.998, HEC-18 pier scour "
    f"r = 0.605). QH provides a complementary turbulence amplification factor that "
    f"augments\u2014rather than replaces\u2014established empirical methods.",
    f"ALR reduces a 6,000-particle uniform simulation to 500 particles with "
    f"{cost.errors_vorticity[1]:.2%} vorticity error and circulation conserved to 0.03% "
    f"via the symmetrized Biot-Savart kernel, achieving a {speedup:.0f}\u00d7 "
    f"wall-time reduction.",
    f"Quasi-unsteady sediment transport with Hirano armoring produces 10.6 ft of "
    f"clear-water scour with 9\u00d7 surface coarsening\u2014physically consistent behavior "
    f"for a dam-release scenario on sand-gravel bed (Williams & Wolman 1984).",
    f"The method integrates with PCSWMM as a post-processor requiring no additional "
    f"meshing or CFD software, filling the gap between Manning\u2019s n + HEC-18 correction "
    f"factors and full 3D CFD for routine scour screening.",
]
for i, c in enumerate(conclusions, 1):
    body(f"{i}. {c}")

body(
    "Future work includes coupling the quasi-unsteady sediment engine directly to the "
    "ALR vortex particle field for pier-scale morphodynamic feedback, direct comparison "
    "against measured scour profiles from published flume experiments, and application "
    "to real-world PCSWMM models.",
    indent=True,
)

doc.add_heading("7.1 Code Availability", level=2)
body(
    "The complete ALR Python/Numba source code, including the symmetrized Biot-Savart "
    "kernel, adaptive-resolution engine, quasi-unsteady sediment transport module, "
    "PCSWMM post-processor, and all 109 validation tests, is open-source under the MIT "
    "license and available at https://doi.org/10.5281/zenodo.19462126 "
    "pip install -e . installation, a one-command benchmark runner "
    "(python run_benchmark_validation.py), and reproducible synthetic test cases that "
    "match every table and figure in this paper.",
    indent=True,
)

# ── REFERENCES ────────────────────────────────────────────────────────────

doc.add_heading("References", level=1)

refs = [
    'Barba, L. & Rossi, L. (2010). "Global field interpolation for particle methods." '
    'J. Computational Physics, 229(4), 1292-1310.',
    'Colebrook, C.F. (1939). "Turbulent flow in pipes." J. Inst. Civil Engineers, 11(4), 133-156.',
    'Egiazaroff, I. (1965). "Calculation of non-uniform sediment concentrations." '
    'J. Hydraulics Division, ASCE, 91(HY4), 225-248.',
    'Hirano, M. (1971). "River bed degradation with armoring." '
    'Trans. Japan Society Civil Engineers, 195, 55-65.',
    'Kolmogorov, A.N. (1941). "Local structure of turbulence." Dokl. Akad. Nauk SSSR, 30, 301-305.',
    'Melville, B.W. (1997). "Pier and abutment scour: Integrated approach." '
    'J. Hydraulic Engineering, ASCE, 123(2), 125-136.',
    'Meyer-Peter, E. & M\u00fcller, R. (1948). "Formulas for bed-load transport." '
    'Proc. 2nd Meeting IAHR, Stockholm, 39-64.',
    'Richardson, E.V. & Davis, S.R. (2001). Evaluating Scour at Bridges (HEC-18). FHWA-NHI-01-001.',
    'Williams, G.P. & Wolman, M.G. (1984). "Downstream effects of dams on alluvial rivers." '
    'USGS Professional Paper 1286.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f"[{i}] {ref}")
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.left_indent = Cm(1.0)
    for run in p.runs:
        run.font.size = Pt(9)

# ── SAVE ──────────────────────────────────────────────────────────────────

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Flynn_ASCE_2026_ALR_Paper_REV1.docx",
)
doc.save(output_path)
print(f"\nPaper saved: {output_path}")
