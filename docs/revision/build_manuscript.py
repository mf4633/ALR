"""Build a clean revised manuscript by editing the original docx in place,
preserving equations and formatting. Swaps affected paragraphs and tables."""
import docx, copy

SRC = '/root/.claude/uploads/03a42327-6a18-5998-96ca-a9e2c0a797b2/cda7807f-Flynn_JWMM_2026_ALR_Paper.docx'
OUT = '/tmp/claude-0/-home-user-ALR/03a42327-6a18-5998-96ca-a9e2c0a797b2/scratchpad/Flynn_JWMM_2026_ALR_Paper_REVISED.docx'

d = docx.Document(SRC)

def set_para_text(p, new_text):
    """Replace all text in a paragraph, preserving its style & first run format."""
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)

# (unique substring in the ORIGINAL paragraph)  ->  full clean replacement text
REPLACE = {
 # ---- Abstract ----
 "Five controlled experiments show that a 500-particle run recovers":
 ("Hydraulic engineers in practice face a pronounced modeling gap. The Manning equation is fast and inexpensive but collapses the turbulent boundary layer into a single empirical roughness coefficient. Three-dimensional computational fluid dynamics resolves that boundary layer but typically requires weeks of setup, specialized software, and a dedicated consultant, placing it out of reach for preliminary scour screening, design-alternative comparison, and routine stormwater, transportation, and dam-safety work. This paper presents a vortex-particle post-processor for the Storm Water Management Model and its Personal Computer variant that provides physics-based turbulence information at user-selected critical locations without mesh generation or outside expertise. The method, termed Adaptive Lagrangian Refinement, uses an observation-dependent resolution rule: particle core sizes contract near locations the engineer designates as critical and expand elsewhere, concentrating computation only where it affects the answer. The underlying symmetrized variable-blob Biot–Savart kernel preserves circulation to 0.03%. Because vorticity error scales as one over the square root of the particle count, observation-concentrated placement reaches a given in-zone accuracy with approximately 4.7 times fewer particles than uniform placement, at correspondingly lower wall time. Benchmark validation cross-checks the Tier-1 bed-shear field against six empirical design families for bridge-pier and contraction scour and against published computational-fluid-dynamics results from the bridge-pier literature. The reference implementation is open-source."),

 # ---- 1.3 contribution 1 ----
 "termed Adaptive Lagrangian Refinement (ALR) in what follows, that concentrates":
 ("1. An observation-dependent resolution rule for vortex-particle methods, termed Adaptive Lagrangian Refinement (ALR) in what follows, that concentrates computation only where the engineer requests it, reaching a given in-zone accuracy with roughly 4.7 times fewer particles than uniform placement."),

 # ---- 3.4 Melville amplification paragraph ----
 "For this configuration (3-ft / 0.9-m pier in a 40-ft / 12-m channel, blockage ratio 0.075)":
 ("For this configuration (3-ft / 0.9-m pier in a 40-ft / 12-m channel, blockage ratio 0.075), the Tier-1 Colebrook–White analysis provides a geometric-blockage shear amplification of ~1.11×, independent of flow intensity (Table 3). This geometric factor is complementary to Melville’s flow-intensity parameter K_I: the Tier-1 analysis captures the blockage (pier width / channel width), while Melville’s K_I captures the flow-intensity dependence. The two factors are complementary—the geometric amplification augments the Melville baseline, it does not replace it. A practitioner would use both:"),

 # ---- 3.5 the ALR/CFD positioning paragraph (starts "ALR’s Tier-2 ...") ----
 "ALR’s Tier-2 vortex-particle analysis at the synthetic 3-ft (0.9-m) circular pier":
 ("ALR does not attempt to reproduce this near-pier peak. Its Tier-1 analysis provides a zone-scale geometric-blockage amplification (~1.09× at the synthetic 3-ft / 0.9-m pier of Section 4.3), and its Tier-2 vortex-particle step, seeded from the mean shear over an engineering-scale observation zone, correctly adds no further bed-shear amplification: a mean-shear-only vortex field satisfies the constant-stress-layer limit (minus the Reynolds shear stress tends to the square of the friction velocity) and therefore returns unit amplification. Resolving the CFD-scale peak would require injecting the coherent near-pier structures (horseshoe and shedding vortices) into the control volume and calibrating their magnitude against flume or CFD data; this is identified as future work (Section 5), not a result claimed here."),

 # ---- Table 4 caption ----
 "Positioning of ALR between Manning’s + HEC-18 screening and full 3D RANS":
 ("Table 4 Positioning of ALR between Manning’s + HEC-18 screening and full 3D RANS (Roulund et al., 2005) for a circular-pier bed-shear analysis. Peak shear values from Roulund are reported in τ/τ∞ at the boundary-layer separation line; ALR reports a Tier-1 geometric-blockage amplification over the engineer-defined observation zone, with the mean-shear Tier-2 step adding no further amplification."),

 # ---- 4.2 performance ----
 "At the optimal operating point of 500 ALR particles, the method achieves a 19×":
 ("Because vorticity error scales as one over the square root of the particle count, the natural cost–benefit metric is the particle count required to reach a given in-zone accuracy. Across particle budgets from 500 to 16,000, observation-concentrated placement reaches the same in-zone vorticity accuracy as uniform placement with approximately 4.7 times fewer particles (Table 6), at correspondingly lower wall time under the sub-quadratic 6σ-cutoff scaling (Section 5.4). This in-zone benefit is accompanied, as expected, by higher error outside the observation zone—the resolution the method deliberately removes from regions the engineer has not designated as critical."),

 # ---- 4.3 engineering scour ----
 "At a synthetic bridge pier, Tier 2 (vortex particles) amplifies bed shear by 1.44×":
 ("At the synthetic 3-ft (0.9-m) bridge pier, Tier 1 (vectorized Colebrook–White) reports a geometric-blockage bed-shear amplification of ~1.09× (pier-zone 0.138 psf versus approach 0.126 psf). The Tier 2 vortex-particle step, seeded from the mean-shear vorticity over the observation zone, returns an amplification of 1.00×—it adds no bed-shear amplification beyond the Tier-1 geometric estimate, consistent with the constant-stress-layer limit. The reported Shields parameter at the pier is 0.86 and the logistic scour-severity index is 0.88. The physical content of the Tier-2 step at this configuration is the resolved turbulence-intensity / turbulent-kinetic-energy field, not a shear multiplier."),

 # ---- Figure 5 caption ----
 "Surface d50 coarsening from 0.80":
 ("Figure 5 Surface d50 coarsening from 0.80 mm (medium sand) to 5.0 mm (fine gravel) over the hydrograph, illustrating the Hirano active-layer armoring mechanism (mass-conserving engine)."),

 # ---- 4.4 sediment result ----
 "increase that forms a gravel armor and limits further transport":
 ("The surface coarsened from 0.80 mm (medium sand) to 5.0 mm (fine gravel)—a 6.3× increase—forming a gravel armor that limits further transport, with the bed degrading by 0.31 ft. Mass is conserved to about 1 part in 10^15 over the hydrograph, and the erosion is transport-limited (the per-step Courant limiter never binds). This self-limiting behavior is the primary physical mechanism controlling long-term degradation below dams and is consistent with field observations (Williams and Wolman, 1984). The demonstration is included as evidence of extensibility; full pier-scale coupling between the ALR vortex field and the morphodynamic engine is future work (Section 5)."),

 # ---- 5.1 positioning ----
 "The benchmarks of Section 3 show the two faces of that positioning":
 ("ALR occupies the middle ground. The benchmarks of Section 3 show the two faces of that positioning: good correlation with empirical scour formulas via the Tier-1 geometric bed-shear field (Laursen r = 0.998, HEC-18 r = 0.605; Sections 3.2–3.3), together with a resolved turbulence-intensity diagnostic at designated locations. The computational cost is in the same order of magnitude as Manning’s (0.021 s per pier scenario on a standard laptop), rather than the hours-to-days of CFD."),

 # ---- 5.4 scaling ----
 "The Biot–Savart evaluation uses a  cutoff radius that limits pairwise interaction":
 ("The Biot–Savart evaluation uses a 6σ cutoff radius that limits pairwise interaction to O(N·K), where K is the average neighbor count, rather than O(N²). Larger core sizes in coarse zones further reduce K. Because vorticity error scales as one over the square root of the particle count, observation-concentrated placement reaches a given in-zone accuracy with roughly 4.7 times fewer particles than uniform placement, with correspondingly lower wall time. The implementation uses Numba just-in-time compilation for the inner loop. All five ALR experiments plus the six benchmark suites complete in under one minute on a standard laptop."),

 # ---- Conclusion 2 ----
 "The core algorithmic contribution is an observation-dependent resolution rule implemented on a symmetrized":
 ("2. The core algorithmic contribution is an observation-dependent resolution rule implemented on a symmetrized Biot–Savart kernel (Barba and Rossi, 2010). Observation-concentrated placement reaches a given in-zone vorticity accuracy with approximately 4.7 times fewer particles than uniform placement, with circulation conserved to 0.03%."),

 # ---- Conclusion 4 ----
 "ALR is not a substitute for CFD where peak shear at the separation line is the controlling quantity":
 ("4. Comparison against the published CFD dataset of Roulund et al. (2005) situates ALR between Manning’s + HEC-18 screening and boundary-layer-resolved 3D RANS. ALR provides a Tier-1 geometric bed-shear field and a resolved turbulence-intensity diagnostic; it does not reproduce the near-pier CFD peak, which would require injected coherent structures with calibrated magnitude (future work). It is an upgrade over Manning’s for the much larger class of routine work where CFD is economically unavailable."),

 # ---- Conclusion 5 ----
 "Extensibility of the framework is demonstrated by coupling the ALR particle field to a quasi-unsteady":
 ("5. Extensibility of the framework is demonstrated by coupling the ALR particle field to a mass-conserving quasi-unsteady fractional sediment-transport engine with Hirano active-layer armoring, producing self-limiting surface coarsening (0.80 to 5.0 mm, 6.3×) and 0.31 ft of bed degradation for a dam-release scenario consistent with field observations (Williams and Wolman, 1984)."),
}

n_par = 0
for p in d.paragraphs:
    for key, new in list(REPLACE.items()):
        if key in p.text:
            set_para_text(p, new)
            del REPLACE[key]
            n_par += 1
            break
print("paragraphs replaced:", n_par)
if REPLACE:
    print("NOT MATCHED:")
    for k in REPLACE: print("  -", k[:60])

# ---------- Tables ----------
def set_cell(t, r, c, val):
    t.rows[r].cells[c].text = val

# Table index 3 = positioning table: fix the near-pier row (1.44x -> 1.00x)
t3 = d.tables[3]
for row in t3.rows:
    if "Near-pier" in row.cells[0].text:
        row.cells[2].text = "1.09× geometric (Tier 1); Tier 2 mean-shear adds 1.00× (turbulence-intensity diagnostic)"
print("Table 3 near-pier row updated")

# Table index 4 = convergence: refresh to current converged values
conv = [("5","0.6929","0.34792","0.18532"),
        ("10","0.4902","0.34792","0.18532"),
        ("15","0.3329","0.34792","0.18532"),
        ("25","0.2153","0.34769","0.18523"),
        ("50","0.1722","0.34543","0.18218"),
        ("100","0.1630","0.34543","0.18218")]
t4 = d.tables[4]
for i,(rad,sig,vor,ens) in enumerate(conv):
    r = i+1
    set_cell(t4,r,0,rad); set_cell(t4,r,1,sig); set_cell(t4,r,2,vor); set_cell(t4,r,3,ens)
print("Table 4 convergence refreshed")

# Table index 5 = performance -> honest cost-benefit (rebuild rows)
t5 = d.tables[5]
hdr = ["N particles","Wall time (s)","In-zone err (uniform)","In-zone err (concentrated)"]
rows = [("500","0.027","1.162","0.626"),
        ("1,000","0.040","1.090","0.432"),
        ("2,000","0.058","0.659","0.369"),
        ("4,000","0.094","0.505","0.226"),
        ("8,000","0.161","0.358","0.145"),
        ("16,000","0.389","0.202","0.118")]
for c,h in enumerate(hdr): set_cell(t5,0,c,h)
for i,row in enumerate(rows):
    for c,v in enumerate(row): set_cell(t5,i+1,c,v)
print("Table 5 -> cost-benefit (~4.7x particle reduction to equal in-zone accuracy)")

# Table index 6 = sediment metrics
t6 = d.tables[6]
sed = [("Metric","Value"),
       ("Total bed change","-0.31 ft"),
       ("Initial surface d50","0.80 mm"),
       ("Final surface d50","5.0 mm"),
       ("Coarsening ratio","6.3×"),
       ("Armor formed","Yes")]
for i,(k,v) in enumerate(sed):
    set_cell(t6,i,0,k); set_cell(t6,i,1,v)
# add mass-conservation row if space; else note in metric
print("Table 6 sediment refreshed")

d.save(OUT)
print("SAVED", OUT)
